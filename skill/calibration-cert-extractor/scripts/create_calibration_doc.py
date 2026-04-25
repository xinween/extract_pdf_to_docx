#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
使用docx库创建仪表清单文档，保留模板样式
用法: python create_calibration_doc.py <template.docx> <output.docx> [data.json]

关键说明：
  - 表格结构：行0 = 表头第一层（Calibration 校准 跨列合并），行1 = 表头第二层（Cal. Date / Due Date）
  - 数据从行2开始
  - 本脚本通过复制已有数据行的 XML 来保留模板样式，不调用 add_row()
  - 动态列映射：通过读取表头文字自动识别每列用途，支持不同列数的模板
  - 序号列自动检测：若表头行对应列已含序号数字则跳过不写入（保留模板自动编号）
"""

import sys
import json
import copy
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_ROWS = 2  # 模板表头占 2 行（行0: Calibration 合并行，行1: Cal.Date/Due Date 分行）


# ── 列名关键词映射表 ───────────────────────────────────────────────────────────
# key: 数据字段名  value: 可能出现在表头中的关键词列表（不区分大小写）
COLUMN_KEYWORDS = {
    'no':       ['no.', '序号', 'no |', 'number'],
    'name':     ['instruments name', 'instrument name', '仪器仪表名称', '仪器名称'],
    'pid':      ['p&id', 'p &id', 'pid'],
    'sn':       ['series no', 'serial no', '序列号'],
    'accuracy': ['accuracy', 'aaccuracy', '精度'],       # 部分模板有此列
    'range':    ['range', '量程'],                        # 部分模板有此列
    'brand':    ['brand', '品牌'],
    'model':    ['model', '型号'],
    'cert':     ['calibration certificate', 'certificate no', '校准证书编号', '证书编号'],
    'cal_date': ['cal. date', 'cal date', '校准日期'],
    'due':      ['due date', '有效期'],
    'remark':   ['remark', '备注'],
}


def detect_column_map(table, header_rows=2):
    """
    通过扫描表头行，自动生成列索引映射。
    返回: dict {字段名: 列索引}，未匹配的字段不在 dict 中。

    策略：
    - 扫描所有表头行（前 header_rows 行）的单元格文字
    - 对每个字段，找到第一个包含其关键词的列
    - 'cal_date' 和 'due' 同在"Calibration 校准"合并单元格下方，通过行1区分
    """
    n_cols = len(table.columns)
    col_map = {}

    # 收集每列的全部表头文字（多行拼接，小写）
    col_texts = []
    for c in range(n_cols):
        parts = []
        for r in range(min(header_rows, len(table.rows))):
            try:
                txt = table.rows[r].cells[c].text.strip().lower()
                if txt:
                    parts.append(txt)
            except Exception:
                pass
        col_texts.append(' | '.join(parts))

    # 对每个字段关键词，扫描找到列索引
    for field, kws in COLUMN_KEYWORDS.items():
        for c, txt in enumerate(col_texts):
            for kw in kws:
                if kw.lower() in txt:
                    if field not in col_map:   # 只取第一个匹配
                        col_map[field] = c
                    break
            if field in col_map:
                break

    return col_map


def is_auto_numbered(table, col_idx, header_rows=2):
    """
    检测某列是否为 Word 自动编号（不需要脚本写入序号）。

    两种情形均视为"自动编号"：
    1. Word 列表编号域（w:numId）：单元格 text 为空，但含有 numId XML 节点
    2. 模板已预置连续数字（1,2,3...）：直接扫描文字内容
    """
    data_rows = list(table.rows[header_rows:])
    if not data_rows:
        return False

    # ── 情形1：检测 Word 列表编号（w:numId） ──────────────────────────────
    num_id_count = 0
    for row in data_rows[:5]:  # 只检查前5个数据行
        try:
            tc = row.cells[col_idx]._tc
            if tc.find('.//' + qn('w:numId')) is not None:
                num_id_count += 1
        except Exception:
            pass
    if num_id_count >= 1:   # 只要有一行检测到 numId 就认定为自动编号
        return True

    # ── 情形2：预置连续数字文本 ───────────────────────────────────────────
    numbers_found = []
    for row in data_rows:
        try:
            txt = row.cells[col_idx].text.strip()
            if txt.isdigit():
                numbers_found.append(int(txt))
        except Exception:
            pass
    if len(numbers_found) >= 2:
        numbers_found.sort()
        return numbers_found[0] == 1 and all(
            numbers_found[i+1] == numbers_found[i]+1
            for i in range(len(numbers_found)-1)
        )

    return False


def set_vertical_align(cell, align='center'):
    """设置单元格垂直对齐: top, center, bottom"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    vAlign = OxmlElement('w:vAlign')
    val_map = {'center': 'center', 'bottom': 'bottom', 'top': 'top'}
    vAlign.set(qn('w:val'), val_map.get(align, 'center'))
    tcPr.append(vAlign)


def clear_cell_text(cell):
    """清空单元格所有段落文本，但保留段落对象本身（保留样式）"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
        runs = para.runs
        for run in runs[1:]:
            run._r.getparent().remove(run._r)


def set_cell_text(cell, text, font_size=9, center=True):
    """设置单元格文本，尽量复用模板已有段落/run 样式"""
    for i, para in enumerate(cell.paragraphs):
        if i == 0:
            for run in para.runs:
                run.text = ''
            runs = para.runs
            for run in runs[1:]:
                run._r.getparent().remove(run._r)
            if para.runs:
                para.runs[0].text = text
                para.runs[0].font.size = Pt(font_size)
            else:
                run = para.add_run(text)
                run.font.size = Pt(font_size)
            if center:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p_elem = para._p
            p_elem.getparent().remove(p_elem)

    set_vertical_align(cell, 'center')


def set_bilingual_cell(cell, en_text, cn_text, font_size=9, center=True):
    """设置双语单元格（英文在上，中文在下），复用模板样式"""
    combined = en_text + '\n' + cn_text if en_text and cn_text else (en_text or cn_text)

    for i, para in enumerate(cell.paragraphs):
        if i > 0:
            p_elem = para._p
            p_elem.getparent().remove(p_elem)

    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ''
    runs = para.runs
    for run in runs[1:]:
        run._r.getparent().remove(run._r)

    if para.runs:
        para.runs[0].text = combined
        para.runs[0].font.size = Pt(font_size)
    else:
        run = para.add_run(combined)
        run.font.size = Pt(font_size)

    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_vertical_align(cell, 'center')


def clone_row(template_row, data_table):
    """
    深拷贝一个模板行，插入到表格末尾，并清空所有单元格内容。
    返回新行对象。
    """
    tbl = data_table._tbl
    new_tr = copy.deepcopy(template_row._tr)
    tbl.append(new_tr)

    from docx.table import _Row
    new_row = _Row(new_tr, data_table)

    for cell in new_row.cells:
        clear_cell_text(cell)

    return new_row


def create_calibration_doc(template_path, output_path, data):
    """
    创建校准证书清单文档

    Args:
        template_path: 模板文件路径
        output_path:   输出文件路径
        data:          列表，每项为 dict，字段见 SKILL.md
    """
    doc = Document(template_path)

    tables = doc.tables
    if len(tables) < 4:
        raise ValueError(f"模板中只有 {len(tables)} 个表格，需要至少4个表格")

    data_table = tables[3]
    tbl = data_table._tbl

    # ── 步骤0：动态识别列映射 ─────────────────────────────────────────────────
    col_map = detect_column_map(data_table, HEADER_ROWS)
    print(f"[列映射] 检测结果: {col_map}")

    # 检测序号列是否为自动编号
    no_col = col_map.get('no')
    skip_no = False
    if no_col is not None:
        skip_no = is_auto_numbered(data_table, no_col, HEADER_ROWS)
        print(f"[序号列] col={no_col}, 自动编号={skip_no} → {'跳过写入' if skip_no else '正常写入'}")

    # ── 步骤1：保留表头行（前 HEADER_ROWS 行），删除其余行 ──────────────────
    all_rows = data_table.rows
    rows_to_remove = list(all_rows[HEADER_ROWS:])

    template_data_row = rows_to_remove[0] if rows_to_remove else None

    for row in rows_to_remove:
        tbl.remove(row._tr)

    if template_data_row is None:
        template_data_row = data_table.rows[-1]

    # ── 步骤2：逐条插入数据行 ────────────────────────────────────────────────
    for item in data:
        row = clone_row(template_data_row, data_table)
        cells = row.cells

        def write(field, text, bilingual=False, en='', cn=''):
            """往 col_map 中找到的列写入数据，找不到则跳过"""
            cidx = col_map.get(field)
            if cidx is None or cidx >= len(cells):
                return
            if bilingual:
                if en and cn:
                    set_bilingual_cell(cells[cidx], en, cn)
                else:
                    set_cell_text(cells[cidx], en or cn)
            else:
                set_cell_text(cells[cidx], text)

        # 序号（自动编号时跳过）
        if not skip_no:
            write('no', str(item.get('no', '')))

        # 仪器仪表名称（双语）
        write('name', '',
              bilingual=True,
              en=item.get('name_en', ''),
              cn=item.get('name_cn', ''))

        # P&ID 编号
        write('pid', item.get('pid', ''))

        # 序列号
        write('sn', item.get('sn', ''))

        # 精度（仅部分模板有此列）
        write('accuracy', item.get('accuracy', ''))

        # 量程（仅部分模板有此列）
        write('range', item.get('range', ''))

        # 品牌（双语）
        write('brand', '',
              bilingual=True,
              en=item.get('brand_en', ''),
              cn=item.get('brand_cn', ''))

        # 型号
        write('model', item.get('model', ''))

        # 校准证书编号
        write('cert', item.get('cert', ''))

        # 校准日期
        write('cal_date', item.get('cal_date', ''))

        # 有效期至
        write('due', item.get('due', ''))

        # 备注
        write('remark', item.get('remark', ''))

    doc.save(output_path)
    print(f"文档已保存: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("用法: python create_calibration_doc.py <template.docx> <output.docx> [data.json]")
        sys.exit(1)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    data_path = sys.argv[3] if len(sys.argv) > 3 else None

    if data_path:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        json_data = input("请输入JSON数据: ")
        data = json.loads(json_data)

    create_calibration_doc(template_path, output_path, data)


if __name__ == '__main__':
    main()
